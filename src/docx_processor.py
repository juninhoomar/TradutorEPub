import time
from docx import Document
from docx.shared import Pt
from src.translator import Translator
from tqdm import tqdm
import concurrent.futures
import re

class DocxProcessor:
    def __init__(self, input_file, output_file, progress_callback=None):
        self.input_file = input_file
        self.output_file = output_file
        self.progress_callback = progress_callback
        self.translator = Translator()
        self.doc = Document(input_file)

    def process(self):
        """
        Main processing loop.
        """
        print(f"Starting translation of {self.input_file}...")
        
        # Collect all paragraphs that need translation
        paragraphs_to_translate = []
        for p in self.doc.paragraphs:
            if p.text.strip():
                paragraphs_to_translate.append(p)
        
        # Also collect tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.strip():
                            paragraphs_to_translate.append(p)

        total_paragraphs = len(paragraphs_to_translate)
        print(f"Found {total_paragraphs} paragraphs to translate.")

        # We can parallelize the translation requests
        # However, we need to keep track of which result belongs to which paragraph
        # And we need to be careful with rate limits.
        
        # Let's process in chunks to show progress and manage memory/errors
        # chunk_size = 10 # Adjust based on API limits
        
        # Group paragraphs into blocks for efficient translation
        blocks = []
        current_block = []
        current_char_count = 0
        MAX_BLOCK_CHARS = 2000
        MAX_BLOCK_PARAGRAPHS = 10

        for p in paragraphs_to_translate:
            # Check if paragraph is "complex" (multiple runs with different formatting)
            is_complex = len(p.runs) > 1
            
            if is_complex:
                if current_block:
                    blocks.append(current_block)
                    current_block = []
                    current_char_count = 0
                blocks.append(p)
            else:
                text_len = len(p.text)
                if (len(current_block) >= MAX_BLOCK_PARAGRAPHS or 
                    current_char_count + text_len > MAX_BLOCK_CHARS):
                    blocks.append(current_block)
                    current_block = []
                    current_char_count = 0
                
                current_block.append(p)
                current_char_count += text_len
        
        if current_block:
            blocks.append(current_block)

        total_tasks = len(blocks)
        print(f"Created {total_tasks} translation tasks (blocks + complex paragraphs).")

        # Use ThreadPoolExecutor for concurrent API calls
        with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
            futures = []
            for item in blocks:
                if isinstance(item, list):
                    futures.append(executor.submit(self.process_block, item))
                else:
                    futures.append(executor.submit(self.process_paragraph, item))
            
            # Monitor progress
            completed_count = 0
            for _ in tqdm(concurrent.futures.as_completed(futures), total=len(futures), desc="Translating"):
                completed_count += 1
                if self.progress_callback:
                    self.progress_callback(completed_count, total_tasks, "Traduzindo (Blocos)")
                    
        print("Saving document...")
        self.doc.save(self.output_file)
        print(f"Translation saved to {self.output_file}")


    def process_block(self, block):
        """
        Translates a block of paragraphs as a single unit to maintain context.
        """
        if not block:
            return

        tagged_text = ""
        valid_indices = []
        for i, p in enumerate(block):
            text = p.text.strip()
            if text:
                tagged_text += f"<p id={i}>{text}</p>\n"
                valid_indices.append(i)
        
        if not tagged_text:
            return

        prompt = (
            "Translate the following text from English to Portuguese (Brazil).\n"
            "The text is structured as a block of paragraphs with tags like <p id=N>text</p>.\n"
            "You MUST preserve these tags and their IDs exactly.\n"
            "Translate the content INSIDE the tags.\n\n"
            "Pay attention to:\n"
            "- **False cognates**.\n"
            "- **Gender and number agreement**.\n"
            "- **Natural word order**.\n"
            "- **Idiomatic expressions**.\n"
            "- **Maintain the flow and coherence between paragraphs**.\n\n"
            "Do not reorder the tags.\n"
            "Output ONLY the translated tagged text."
        )

        try:
            translated_tagged = self.translator.translate_text(tagged_text, system_instruction=prompt)
            
            matches = re.finditer(r"<p id=(\d+)>(.*?)</p>", translated_tagged, re.DOTALL)
            
            found_ids = set()
            for match in matches:
                p_id = int(match.group(1))
                content = match.group(2).strip()
                found_ids.add(p_id)
                
                if 0 <= p_id < len(block):
                    p = block[p_id]
                    # Update paragraph text
                    if p.runs:
                        p.runs[0].text = content
                        # Clear other runs if any
                        for k in range(1, len(p.runs)):
                            p.runs[k].text = ""
                    else:
                        p.add_run(content)
            
            # Check for missing paragraphs and fallback
            missing_ids = set(valid_indices) - found_ids
            if missing_ids:
                print(f"Warning: {len(missing_ids)} paragraphs missed in block translation. Falling back to individual translation.")
                for mid in missing_ids:
                    self.process_paragraph(block[mid])

        except Exception as e:
            print(f"Error processing block: {e}. Falling back to individual translation.")
            for p in block:
                self.process_paragraph(p)

    def process_paragraph(self, paragraph):
        """
        Translates a single paragraph preserving run formatting.
        """
        try:
            # check if paragraph has mixed formatting or multiple runs
            if len(paragraph.runs) <= 1:
                # Simple case
                text = paragraph.text
                translated_text = self.translator.translate_text(text)
                if paragraph.runs:
                    paragraph.runs[0].text = translated_text
                else:
                    paragraph.add_run(translated_text)
                return

            # Complex case: multiple runs
            # Construct tagged text
            tagged_text = ""
            for i, run in enumerate(paragraph.runs):
                # Use a marker that is unlikely to appear in text. 
                # Using XML-like tags might confuse the model if it tries to parse them as HTML.
                # Let's use simple IDs.
                tagged_text += f" <run id={i}>{run.text}</run>"
            
            # Translate with instruction to preserve tags
            # We need a specialized prompt for this in the translator, 
            # but for now we'll rely on the generic one or modify the translator call.
            # Actually, let's just pass the tagged text and hope the model respects it.
            # To ensure it respects it, we might need a specific system prompt for this.
            
            # Since we can't easily change the Translator class method signature without refactoring,
            # let's just pass it. The Translator has a generic prompt.
            # We should probably update Translator to handle this or just subclass/modify here.
            # Let's do a direct call here for better control or update Translator.
            # I'll update Translator to be more flexible? 
            # No, let's just try to send it.
            # Better: Update the prompt dynamically.
            
            prompt = """
            Translate the following text from English to Portuguese (Brazil). 
            The text contains tags like <run id=N>text</run>. 
            You MUST preserve these tags and their IDs exactly. 
            Translate the content INSIDE the tags.
            
            Pay attention to:
            - **False cognates** ("pretend" ≠ "pretender", "push" ≠ "puxar", etc.).
            - **Gender and number agreement**.
            - **Natural word order**.
            - **Idiomatic expressions**.
            
            Do not reorder the tags.
            Do not repeat the content.
            If a run contains only whitespace or symbols, keep it as is.
            Output ONLY the translated tagged text.
            """
            
            translated_tagged = self.translator.translate_text(tagged_text, system_instruction=prompt)
            
            # Parse the result
            # We expect <run id=N>...</run>
            # We use regex to find them.
            
            matches = re.finditer(r"<run id=(\d+)>(.*?)</run>", translated_tagged, re.DOTALL)
            
            found_runs = {}
            for match in matches:
                run_id = int(match.group(1))
                content = match.group(2)
                found_runs[run_id] = content
            
            # Validation: If we didn't find most of the runs, something went wrong.
            # It's better to fall back to simple translation than to have mixed languages.
            if len(found_runs) < len(paragraph.runs):
                # Check if the missing runs are just empty/whitespace
                missing_ids = set(range(len(paragraph.runs))) - set(found_runs.keys())
                significant_missing = False
                for mid in missing_ids:
                    if paragraph.runs[mid].text.strip():
                        significant_missing = True
                        break
                
                if significant_missing:
                    raise ValueError(f"Translation dropped significant runs. Found {len(found_runs)}/{len(paragraph.runs)}")

            # Update runs
            for i, run in enumerate(paragraph.runs):
                if i in found_runs:
                    run.text = found_runs[i]
                else:
                    # If it's not in found_runs, and we passed validation, it means it was likely empty/insignificant
                    # or explicitly removed. We should probably clear it to avoid "English leftovers".
                    # But to be safe against "half-translated" paragraphs, clearing it is better than leaving English.
                    run.text = ""

        except Exception as e:
            print(f"Error processing paragraph: {e}")
            # Fallback: just translate raw text and replace first run
            try:
                text = paragraph.text
                translated = self.translator.translate_text(text)
                for run in paragraph.runs:
                    run.text = ""
                paragraph.runs[0].text = translated
            except:
                pass
