from docx import Document

class UnderlineRemover:
    def __init__(self, input_file, output_file):
        self.input_file = input_file
        self.output_file = output_file
        self.doc = Document(input_file)

    def process(self):
        print(f"Starting underline removal for {self.input_file}...")
        
        # Process regular paragraphs
        for p in self.doc.paragraphs:
            self._process_paragraph(p)
            
        # Process tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._process_paragraph(p)
                        
        print("Saving document...")
        self.doc.save(self.output_file)
        print(f"Processed document saved to {self.output_file}")

    def _process_paragraph(self, paragraph):
        # Identify if a run is inside a hyperlink or is styled as one.
        for run in paragraph.runs:
            is_hyperlink = False
            
            # Check by style name
            if run.style and hasattr(run.style, 'name') and run.style.name == 'Hyperlink':
                is_hyperlink = True
                
            # Check XML parent tag
            parent_tag = run._r.getparent().tag
            if 'hyperlink' in parent_tag.lower():
                is_hyperlink = True

            # If it's not a hyperlink, remove the underline
            if not is_hyperlink:
                # Remove underline formatting at the run level
                if run.font.underline:
                    run.font.underline = False
