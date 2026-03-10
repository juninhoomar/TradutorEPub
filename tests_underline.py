from docx import Document
from src.underline_remover import UnderlineRemover
import os
from docx.enum.style import WD_STYLE_TYPE

# Create a test document
doc = Document()
# Create Hyperlink style if it doesn't exist
styles = doc.styles
try:
    style = styles.add_style('Hyperlink', WD_STYLE_TYPE.CHARACTER)
except ValueError:
    style = styles['Hyperlink']

doc.add_paragraph("This is normal text.")

p = doc.add_paragraph()
run = p.add_run("This should be underlined but will be removed.")
run.font.underline = True

p2 = doc.add_paragraph("This has a ")
run2 = p2.add_run("link")
run2.style = "Hyperlink"
# Hyperlinks usually don't have explicit run.font.underline set because the style handles it
p2.add_run(" inside.")

# Let's add an explicit underline to see if it survives on a hyperlink (it shouldn't be touched)
p3 = doc.add_paragraph()
run3 = p3.add_run("Explicitly underline link")
run3.style = "Hyperlink"
run3.font.underline = True

doc.save("test_input.docx")

remover = UnderlineRemover("test_input.docx", "test_output.docx")
remover.process()

# Check results
doc_out = Document("test_output.docx")
assert doc_out.paragraphs[0].runs[0].font.underline in (None, False)

# Second paragraph had underline=True, now should be False or None
assert doc_out.paragraphs[1].runs[0].font.underline in (None, False)

# Third paragraph has link style
assert doc_out.paragraphs[2].runs[1].style.name == "Hyperlink"

# Fourth paragraph had explicit underline and hyperlink
assert doc_out.paragraphs[3].runs[0].font.underline == True
assert doc_out.paragraphs[3].runs[0].style.name == "Hyperlink"

print("All tests passed!")
os.remove("test_input.docx")
os.remove("test_output.docx")
